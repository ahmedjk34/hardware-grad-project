"""Front matter and Chapter 1."""

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor

from docbuild import ACCENT, BODY_FONT, field, runs_into


def title_page(doc):
    def line(text, size, bold=False, space=10, colour=None, caps=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space)
        p.paragraph_format.line_spacing = 1.15
        runs_into(p, text)
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.name = BODY_FONT
            r.bold = bold or r.bold
            r.font.all_caps = caps
            if colour is not None and r.font.color.rgb is None:
                r.font.color.rgb = colour
        return p

    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    line("[[UNIVERSITY NAME]]", 16, bold=True, space=4)
    line("[[FACULTY / COLLEGE]]", 13, space=4)
    line("[[DEPARTMENT]]", 13, space=48)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    from docbuild import borders
    borders(p, colour=str(ACCENT), size="6", kinds=("top",))

    line("Vision-Assisted Cartesian Robotic System", 22, bold=True, space=2, colour=ACCENT)
    line("for 3D Block Construction", 22, bold=True, space=10, colour=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    borders(p, colour=str(ACCENT), size="6", kinds=("top",))

    line("A graduation project report submitted in partial fulfilment of the "
         "requirements for the degree of", 11, space=6)
    line("[[DEGREE, e.g. Bachelor of Science in Computer Engineering]]", 12, bold=True, space=46)

    line("Submitted by", 11, space=8)
    line("Ahmed Taher Gharib", 14, bold=True, space=3)
    line("Mohie Aldeen Amjad Halawa", 14, bold=True, space=3)
    line("Khalil Mahmoud Qanabita", 14, bold=True, space=42)

    line("Supervised by", 11, space=8)
    line("[[SUPERVISOR NAME AND TITLE]]", 14, bold=True, space=54)

    line("[[MONTH]] [[YEAR]]", 12, space=0)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def front_matter(rep):
    doc = rep.doc

    def plain_heading(text):
        p = doc.add_heading(text, level=1)
        p.paragraph_format.space_before = Pt(0)
        return p

    # ---- Acknowledgements ----
    plain_heading("Acknowledgements")
    rep.p(
        "We would like to thank our supervisor, [[SUPERVISOR NAME]], for the guidance and "
        "the patience that carried this project from a sketch of a moving arm to a machine "
        "that actually builds something. The advice that mattered most was rarely about "
        "code: it was about measuring before believing, and about being honest in writing "
        "about what a rig can and cannot do.")
    rep.p(
        "We also thank the staff of the [[DEPARTMENT / LAB NAME]] for workshop access, for the "
        "3D printing time this project consumed in quantity, and for tolerating a machine that "
        "ran forty-second cycles in a shared space for weeks. Finally we thank our families and "
        "our classmates, who listened to a great deal of talk about limit switches.")
    rep.p(
        "This project was built as a team. The mechanical build, the electronics, the two "
        "firmwares, the vision pipeline and the browser software were all shared work, and "
        "no part of the system belongs to one of us alone.")

    # ---- Abstract ----
    plain_heading("Abstract")
    rep.p(
        "Small-scale automated assembly is one of those problems that looks solved from a "
        "distance and stops looking solved the moment a real machine has to do it. A gantry "
        "that moves to a coordinate is straightforward. A machine that takes a structure a "
        "person designed, feeds itself one part at a time, picks each part up, turns it the "
        "right way round, stacks it on top of the parts already placed, and knows at every "
        "moment whether that actually happened, is a different problem, and most of the "
        "difference is feedback rather than motion.")
    rep.p(
        "This is where our project comes in: a vision-assisted Cartesian robotic cell that "
        "builds three-dimensional structures out of 2.2 x 6.0 x 1.5 cm wooden blocks with no "
        "human placing a single block by hand. The system is built around three controllers "
        "and one brain. A Raspberry Pi 5 is the sole master: it runs the camera, the web "
        "server, the orchestration and every safety rule. An Arduino MEGA 2560 runs the "
        "gantry, a CoreXY X/Y stage with an added Z axis, a servo claw and a 28BYJ-48 "
        "rotation stepper, and executes a fourteen-phase pick-place-park cycle for every "
        "block. An Arduino Uno runs a separate feeder module, a two-stage hopper gate, a "
        "belt and an alignment servo, and confirms with an exit HC-SR04 and a stage IR sensor that "
        "exactly one block left the hopper and arrived at the fixed pickup point. The two "
        "Arduinos never talk to each other; the Pi is the only thing that couples them, and "
        "it does so over two independent USB serial links.")
    rep.p(
        "The block can be laid either way round, which is why the machine carries two "
        "separately calibrated grids and not one: a vertical grid of 7 x 6 addressable "
        "cells and a horizontal grid of 3 x 10, sharing one 22.8 x 38.0 cm holder-travel "
        "envelope and one feeder cell at [0,0]. An overhead 160-degree fisheye camera watches "
        "the build surface through a colour-corrected, lens-corrected pipeline, detects the "
        "wooden blocks by warm-colour segmentation, and maps camera pixels to physical "
        "centimetres to grid cells through a saved workspace homography.")
    rep.p(
        "The operator side is a browser application served from the Pi: a click-to-build "
        "console over the live camera image, and a 3D Build Studio in which a person designs "
        "a structure, gets immediate physics feedback on support, collisions and toppling, "
        "compiles the design into an ordered program of `B col row level` commands separated "
        "by grid-mode latches, and runs it while a live digital twin mirrors the real build "
        "from the firmware's own phase telemetry.")
    rep.p(
        "The completed system placed blocks on hardware with a measured cycle time of 19.8 to "
        "32.2 seconds per block, with the fixed phases of the cycle repeating to within 0.01 s "
        "across sixteen logged placements, and stacked a five-level tower on a single cell. "
        "The software is covered by 492 browser tests and 68 Python protocol tests. The work "
        "demonstrates that the useful engineering in a machine like this is not the "
        "kinematics: it is the discipline of never letting one part of the system assume "
        "what another part has actually done.")

    # ---- TOC / LoF / LoT ----
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    plain_heading("Table of Contents")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    field(p, r'TOC \o "1-3" \h \z \u',
          placeholder="Right-click here and choose Update Field to build the table of contents.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    plain_heading("List of Figures")
    p = doc.add_paragraph()
    field(p, r'TOC \h \z \c "Figure"',
          placeholder="Right-click here and choose Update Field to build the list of figures.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    plain_heading("List of Tables")
    p = doc.add_paragraph()
    field(p, r'TOC \h \z \c "Table"',
          placeholder="Right-click here and choose Update Field to build the list of tables.")


def chapter_1(rep):
    rep.h1("1. Introduction", page_break=False)

    rep.h2("1.1 Background and Motivation")
    rep.p(
        "Automated pick-and-place is one of the oldest jobs in industrial robotics and it is "
        "still one of the most common: a machine takes a part from a known place, carries it to "
        "a computed place, and lets go. The kinematics of that are well understood, and a "
        "student can put a two-axis gantry together from 3D-printer parts in a few weeks. What "
        "is much less well understood, and much less often built at student scale, is the layer "
        "on top of it. A real assembly cell has to solve where the next part comes from, how it "
        "is presented to the tool in a known orientation, what happens when it is not there, how "
        "the machine knows a part actually landed, and what the machine is allowed to do next "
        "when it does not know.")
    rep.p(
        "Most hobby and teaching machines skip that layer entirely, and they skip it in the same "
        "way: with time. The belt runs for two seconds because two seconds was usually enough. "
        "The gripper closes and the code assumes it closed on something. The arm moves to a "
        "coordinate and the program carries on as though it arrived. That works on a bench, in "
        "front of a person who can see the machine and stop it, and it stops working the moment "
        "the machine has to run a sequence of thirty operations where any one of them can fail "
        "quietly. A block that never left the hopper, a block that was gripped by its corner, a "
        "block that a hand moved while the gantry was somewhere else: none of these produce an "
        "error in a timing-based machine. They produce a wrong structure, built confidently.")
    rep.p(
        "We chose a block-stacking gantry because it is the smallest honest version of that "
        "problem. Wooden blocks are cheap, identical, and forgiving of a millimetre of error, so "
        "the project is not about precision machining. At the same time, stacking is genuinely "
        "three-dimensional and genuinely sequential: block seven has to be right or block "
        "eighteen falls off it, and there is no way to fake that with a well-chosen demo. The "
        "result is a system that is playful to look at, a toy in the best sense, and that has to "
        "get the unglamorous parts right in order to work at all.")
    rep.p(
        "The mechanical starting point for the X/Y stage is the Instructables **Automated "
        "Chessboard** by Greg06 (2022), an open, well-documented CoreXY table that drags "
        "magnetised chess pieces around a board with an electromagnet on a moving trolley. We "
        "started from it for three concrete reasons rather than for convenience. First, its "
        "CoreXY arrangement keeps both motors off the moving beam, so the beam stays light and "
        "the working envelope stays large for a given frame, which is exactly the trade a "
        "student build wants when it is going to add a third axis and a gripper to that beam "
        "later. Second, its homing story is honest and cheap: two limit switches, driven into at "
        "every start, defining a repeatable origin without encoders. Third, the reference is "
        "explicit that a chessboard does not need 3D-printer accuracy, and neither does a "
        "block-stacking rig with 1.6 cm gaps between cells, so the tolerance budget transfers.")
    rep.p(
        "What did not transfer is everything above the table. A chessboard is a plane: the "
        "reference never lifts a piece, never turns one, never has to know what is on the board "
        "except through a fixed matrix of 64 reed switches under it, and never has to supply "
        "itself with pieces. Building in three dimensions turned that reference from a finished "
        "machine into a base frame: what we kept from it is the CoreXY X/Y skeleton and its "
        "limit-switch homing, and the Z axis, the claw, the feeder, the camera and all three "
        "controllers were built on top of that.")
    rep.figure("The completed rig: CoreXY gantry, Z axis and claw, feeder module on the left, "
               "and the overhead camera on its wooden support frame.",
               placeholder="Photograph of the finished machine, three-quarter view, with the "
                           "gantry, the feeder and the camera mount all in shot.")

    rep.h2("1.2 Problem Statement")
    rep.p(
        "This project targets one specific problem: **building a human-designed three-dimensional "
        "block structure, exactly as designed, with no human placing a single block by hand, and "
        "with sensor or vision confirmation at every stage instead of open-loop timing.**")
    rep.p(
        "The problem appears the moment the machine has to run unattended for more than one "
        "operation. Stacking is sequential and unforgiving: each placement is the foundation of "
        "the next, so an error is not an isolated bad block, it is a wrong structure from that "
        "point upward. A timing-based machine cannot tell the difference between a successful "
        "and a failed operation, so it cannot stop; it carries on and turns one silent failure "
        "into a pile. As a result, everything downstream of a wrong placement is wasted, and "
        "worse, a machine that continues after an unknown outcome is a machine driving a gripper into a stack "
        "whose real height it does not know.")
    rep.p(
        "The specific ways an open-loop version of this machine fails are worth naming, because "
        "they are what the design has to answer:")
    rep.bullets([
        "**An empty or jammed hopper.** The gate opens, no block comes out, and the belt runs "
        "for its allotted time anyway. The claw then closes on nothing and places nothing, and "
        "the machine believes a block is at that cell for the rest of the build.",
        "**Two blocks instead of one.** A hopper gate that opens far enough for one block often "
        "opens far enough for two. The second block sits on the belt behind the first and turns "
        "up at the pickup point out of sequence.",
        "**A block that arrives late or crooked.** A belt is not a positioning system. A block "
        "that has not finished travelling, or that has arrived turned a few degrees, is picked "
        "by its corner or missed entirely.",
        "**Lost steps.** Open-loop steppers can lose steps against a stall or a fast direction "
        "reversal, and nothing in the motor reports it. If the position counter is trusted "
        "across a whole build, that error accumulates cell by cell.",
        "**A human in the workspace.** A person nudges a block, or takes one back, and the "
        "machine's model of the board is silently wrong.",
        "**A controller reset.** Opening a USB port resets an Arduino. The board comes back "
        "un-homed, in its compiled default grid mode, with no memory of anything, and a "
        "controller that does not notice keeps sending coordinates into a machine that has "
        "forgotten what they mean.",
    ])
    rep.p(
        "The project focuses instead on making every one of those failures visible and "
        "terminal rather than silent, and on making the machine refuse to continue when it "
        "cannot prove what physically happened. It does not attempt to make the machine decide "
        "**what** to build.")

    rep.h2("1.3 Project Objectives")
    rep.p(
        "This work was done to build a complete, self-feeding robotic cell that constructs a "
        "human-designed 3D block structure end to end, and to make every stage of that "
        "construction confirmed by a sensor or by the camera rather than by a timer.")
    rep.p("The main objectives of this project are:")
    rep.numbered([
        "**Provide** a Cartesian gantry that reaches every cell of a calibrated build grid and "
        "places a block on any addressable cell, at any stack level, within the machine's own "
        "travel envelope. Target: all 41 buildable cells of the vertical grid and all 29 of the "
        "horizontal grid reachable, at levels 0 to 16.",
        "**Enable** both block orientations, so a structure can mix blocks standing and lying, "
        "by turning the claw 90 degrees between the pickup and the placement and carrying a "
        "second, separately calibrated grid for the rotated footprint.",
        "**Deliver** a feeder module that doses exactly one block at a time from a hopper to a "
        "fixed pickup point, and that confirms the block both leaving the container and arriving "
        "at the pickup point with independent exit and stage sensors, so that no placement is ever "
        "started on an assumption that a block is there.",
        "**Support** an overhead vision system that detects the wooden blocks on the build "
        "surface and converts camera pixels into physical centimetres and then into grid cells, "
        "so the camera can be used both to calibrate the machine's grid and to verify what "
        "actually landed. Target: detect every block on a full board and map a cell to within "
        "half a block footprint.",
        "**Provide** browser-based operation from a phone or tablet on the local network: a live "
        "camera view, click-to-build on the image, and an explicit two-tap confirmation before "
        "anything moves, so that the machine can be driven without a keyboard next to it.",
        "**Deliver** a 3D design environment in which a person builds the target structure, gets "
        "immediate feedback on support, collision and toppling, and compiles the design into an "
        "ordered, deterministic command program that the rig then executes block by block.",
        "**Enforce** a safety model in which exactly one command runs at a time, never queued; "
        "in which a failure that leaves the machine's physical state unknown locks the session "
        "and requires a person; and in which no software layer is allowed to retry an operation "
        "whose physical outcome it cannot prove.",
        "**Achieve** a repeatable cycle time per block. Target: under 45 seconds per block, "
        "including feeding, for a machine of this size and step rate.",
    ])
    rep.p(
        "Autonomous planning of **which** block goes where is deliberately not an objective. The "
        "human designs the structure; the build itself, once that design is compiled, is "
        "autonomous, and the machine is responsible for executing it exactly and for stopping "
        "when it cannot.")

    rep.h2("1.4 Scope and Boundaries")
    rep.p(
        "The scope of this project is one complete robotic cell and the software chain that "
        "drives it, from a structure drawn in a browser to a block physically on a stack.")
    rep.p("The following are inside the project:")
    rep.bullets([
        "A Cartesian gantry with CoreXY X/Y motion, an added Z axis, a servo-driven mechanical "
        "claw, and a stepper that rotates the claw 90 degrees, giving two block orientations.",
        "A separate feeder module: a two-stage hopper gate, a conveyor belt, an alignment servo "
        "with an exit ultrasonic sensor and a stage IR sensor, presenting one block at a time "
        "at a fixed pickup point.",
        "An overhead fisheye camera with lens correction, software colour correction, block "
        "detection, and camera-to-machine grid calibration.",
        "A browser operator console served from the Pi, with a live camera view and "
        "click-to-build.",
        "A 3D Build Studio: design, physics validation, compilation to a command program, a "
        "live digital twin, and a guarded execution runner with a run report.",
        "Two independent serial protocols, one per Arduino, and the Pi-side orchestration and "
        "safety layers that couple them.",
    ])
    rep.p("The following aspects are considered out of scope:")
    rep.numbered([
        "Autonomous target selection or any form of AI planning. There is no chess engine, no "
        "solver, and no block-detection-driven building; the operator or a compiled design "
        "chooses every target.",
        "A hardwired emergency stop, safety relay, contactor or physical interlock. This is a "
        "real and acknowledged gap and is discussed in Sections 3.6 and 5.7.",
        "Interrupting a motion already in progress. The gantry firmware does not read serial "
        "during a build cycle, so software stop is stop-after-current-block by construction.",
        "Metrological camera calibration. The lens model is tuned by eye against straight edges "
        "and not fitted from a checkerboard, so the image is visually straight but not "
        "measurement-grade.",
        "WebRTC or hardware-encoded video. The camera is streamed as motion JPEG over the local "
        "network.",
        "Real user accounts and multi-operator coordination. A single shared secret on a trusted "
        "LAN is the whole access model.",
        "Closed-loop motor control. Every stepper in the machine is open-loop; there are no "
        "encoders, and a mechanical stall is not detected by the motor itself.",
    ])
    rep.p(
        "The boundary that matters most is the last one, and it is worth stating precisely "
        "because the word closed-loop is used throughout this report. **Closed-loop here means "
        "sensor- and vision-confirmed sequencing, not servo position control.** The machine "
        "never verifies a motor shaft; it verifies that a block left the hopper, that a block "
        "reached the pickup point, that a build phase began and ended, and that a block is "
        "visible on the surface where one was commanded. Each of those is a fact the machine "
        "measures and not assumes, and each of them can refuse permission for the next step.")

    rep.h2("1.5 Significance of the Project")
    rep.p(
        "The contribution of this project is not the gantry. Cartesian gantries are a solved "
        "problem, and ours is a modified version of a published open design. What this project "
        "demonstrates is the layer that student-scale automation almost always leaves out: an "
        "end-to-end chain in which every physical transition is confirmed by something other "
        "than a clock, and in which the machine is designed to stop rather than guess.")
    rep.p(
        "Three parts of that chain are worth calling out. The **two-board handoff** is the "
        "clearest: a placement is one indivisible operation owned by the Pi, and the gantry is "
        "only allowed to move after the feeder has returned one exact terminal message proving "
        "a block is staged. There is no timeout that means success and no progress message that "
        "counts as permission. The **acknowledged phase protocol** is the second: rather than "
        "letting the Pi guess what is happening during a forty-second silence, the firmware "
        "narrates its own fourteen phases on a separate machine-readable channel, including its "
        "own prediction of how long each Z move should take, which turns an opaque wait into "
        "something a user interface and a log can describe honestly. The **design-compile-twin-run "
        "chain** is the third: a person draws a structure, the compiler turns it into a "
        "deterministic ordered program that respects support order and minimises grid-mode "
        "changes, and a digital twin mirrors the real build from the machine's own telemetry "
        "rather than from an animation timer.")
    rep.p(
        "A system like this is useful to anyone building small automated assembly or sorting "
        "cells where parts have to be presented, oriented and stacked: laboratory sample "
        "handling, kit assembly, educational robotics, and any teaching setting that wants to "
        "show students what the difference between a demo and a machine actually is. It is also "
        "directly useful as a teaching artefact in its own right, because every unsafe shortcut "
        "the design refuses to take is documented next to the reason it was refused.")

    rep.h2("1.6 Organization of the Report")
    rep.p(
        "This report is organized as follows. **Chapter 2** sets out the requirements the system "
        "was designed to, the constraints that shaped it, the overall three-controller "
        "architecture, the mechanical and electrical design, and the end-to-end workflow. "
        "**Chapter 3** covers the hardware implementation subsystem by subsystem: the motion "
        "system and its calibration, the end effector, the three controllers and their division "
        "of responsibility, the vision hardware, the sensors, power and safety, and how the "
        "whole thing is integrated. **Chapter 4** carries the control side and the main weight of "
        "the software: the hierarchical control model, motion control and the grid geometry, the "
        "vision-to-coordinate pipeline, the main automatic operating sequence, and the "
        "additional operating modes. **Chapter 5** presents the testing methodology and the "
        "results, including the measured build timings from the rig's own logs and the automated "
        "test coverage, and discusses the challenges and the limitations that remain. "
        "**Chapter 6** concludes and sets out future work. The appendices carry the pin maps, "
        "the bill of materials, the full command protocols and the raw experimental data.")
