"""Word-document scaffolding for the graduation report.

Everything here is presentation: styles, fields, captions, code blocks, tables,
figure placeholders. The report's text lives in report.py.
"""

import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FIGS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "figs")

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
ACCENT = RGBColor(0x00, 0x00, 0x00)      # headings are black
HEADER_SHADE = "AFC6E9"
FIRSTCOL_SHADE = "D7E3F4"
CODE_SHADE = "F2F3F5"
CODE_COLOR = RGBColor(0x14, 0x3C, 0x8C)
BOX_SHADE = "FFF6DA"


# --------------------------------------------------------------------------
# low-level OOXML helpers
# --------------------------------------------------------------------------

def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def field(paragraph, instruction, placeholder="Right-click and choose Update Field."):
    """A real Word field, so the TOC / caption numbering can be refreshed."""
    r = paragraph.add_run()
    r._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    r = paragraph.add_run()
    t = _el("w:instrText", **{"xml:space": "preserve"})
    t.text = " %s " % instruction
    r._r.append(t)
    r = paragraph.add_run()
    r._r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    if placeholder:
        paragraph.add_run(placeholder)
    r = paragraph.add_run()
    r._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))


def shade(element, colour):
    pr = element.get_or_add_tcPr() if element.tag.endswith("}tc") else element.get_or_add_pPr()
    pr.append(_el("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": colour}))


def borders(paragraph, colour="B8BEC6", size="6", kinds=("top", "left", "bottom", "right")):
    pr = paragraph._p.get_or_add_pPr()
    bd = _el("w:pBdr")
    for k in kinds:
        bd.append(_el("w:" + k, **{"w:val": "single", "w:sz": size,
                                   "w:space": "6", "w:color": colour}))
    pr.append(bd)


def keep_with_next(paragraph):
    paragraph._p.get_or_add_pPr().append(_el("w:keepNext", **{"w:val": "1"}))


# --------------------------------------------------------------------------
# document setup
# --------------------------------------------------------------------------

def new_document():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.25
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    sizes = {"Heading 1": 17, "Heading 2": 14, "Heading 3": 12.5, "Caption": 10}
    for name, size in sizes.items():
        st = doc.styles[name]
        st.font.name = BODY_FONT
        # Heading styles inherit the theme's *major* font, which wins over w:ascii.
        # Drop the theme references so the body font actually applies.
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = _el("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rfonts.get(qn(attr)) is not None:
                del rfonts.attrib[qn(attr)]
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(attr), BODY_FONT)
        if name == "Caption":
            st.font.italic = False
            st.font.color.rgb = RGBColor(0x33, 0x3A, 0x42)
            continue
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = False
        st.font.color.rgb = ACCENT
        st.paragraph_format.space_before = Pt(20 if name == "Heading 1" else 14)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)
        sec.header_distance = sec.footer_distance = Cm(1.25)
    return doc


def running_header(section, text):
    p = section.header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.name = BODY_FONT
    r.font.color.rgb = RGBColor(0x5A, 0x62, 0x6A)
    borders(p, colour="C6CCD2", size="4", kinds=("bottom",))


def page_numbers(section, label="Page ", numfmt="PAGE"):
    p = section.footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.size = Pt(9)
    r.font.name = BODY_FONT
    field(p, numfmt, placeholder="1")
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.name = BODY_FONT
        r.font.color.rgb = RGBColor(0x5A, 0x62, 0x6A)


def start_numbering_at(section, value, fmt="decimal"):
    """Restart page numbering for this section (roman front matter, then arabic)."""
    pr = section._sectPr
    for old in pr.findall(qn("w:pgNumType")):
        pr.remove(old)
    pr.append(_el("w:pgNumType", **{"w:start": str(value), "w:fmt": fmt}))


# --------------------------------------------------------------------------
# inline markup
#   **bold**   __italic__   `code`
#   [[VALUE NEEDED: ...]]   a real unknown — bold, highlighted, unmissable
#   {{INFERRED: ...}}       a deduction — small, grey, bracketed, Ctrl-F-able
# --------------------------------------------------------------------------

_TOKENS = re.compile(r"(\[\[.+?\]\]|\{\{.+?\}\}|\*\*.+?\*\*|`[^`]+`|__.+?__)", re.S)


def runs_into(paragraph, text, size=None, colour=None):
    for chunk in _TOKENS.split(text):
        if not chunk:
            continue
        if chunk.startswith("{{"):
            r = paragraph.add_run("[" + chunk[2:-2] + "]")
            r.italic = True
            r.font.size = Pt((size or 12) - 2)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x7A)
            continue
        if chunk.startswith("[["):
            r = paragraph.add_run(chunk[2:-2])
            r.bold = True
            r.font.color.rgb = RGBColor(0x8A, 0x33, 0x00)
            r._r.get_or_add_rPr().append(_el("w:highlight", **{"w:val": "yellow"}))
        elif chunk.startswith("**"):
            # An identifier quoted inside a bold span renders bold either way, so
            # the backticks are simply dropped instead of printing literally.
            r = paragraph.add_run(chunk[2:-2].replace("`", ""))
            r.bold = True
        elif chunk.startswith("__"):
            r = paragraph.add_run(chunk[2:-2])
            r.italic = True
        elif chunk.startswith("`"):
            # An identifier quoted mid-sentence is body text, not a code block:
            # same font as everything around it, distinguished only by weight.
            r = paragraph.add_run(chunk[1:-1].replace("**", ""))
            r.font.name = BODY_FONT
            r.bold = True
        else:
            r = paragraph.add_run(chunk)
        if size and r.font.size is None:
            r.font.size = Pt(size)
        if colour is not None:
            r.font.color.rgb = colour
    return paragraph


# --------------------------------------------------------------------------
# the authoring API
# --------------------------------------------------------------------------

class Report:
    def __init__(self, doc):
        self.doc = doc
        self.figures = []
        self.tables = []

    # ---- headings ----
    def h1(self, text, page_break=True):
        if page_break:
            self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = self.doc.add_heading(text, level=1)
        return p

    def h2(self, text):
        return self.doc.add_heading(text, level=2)

    def h3(self, text):
        return self.doc.add_heading(text, level=3)

    # ---- body ----
    def p(self, text):
        return runs_into(self.doc.add_paragraph(), text)

    def lead(self, text):
        p = runs_into(self.doc.add_paragraph(), text)
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = RGBColor(0x3A, 0x42, 0x4A)
        p.paragraph_format.space_after = Pt(12)
        return p

    def bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.9)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            runs_into(p, it)

    def numbered(self, items):
        # Numbered explicitly rather than through Word's List Number style, so that
        # every list in the document restarts at 1 instead of continuing the last one.
        for i, it in enumerate(items, 1):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(3)
            pf.left_indent = Cm(1.5)
            pf.first_line_indent = Cm(-0.7)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run("%d.\t" % i)
            r.bold = True
            runs_into(p, it)

    def defs(self, pairs):
        """`Term:` explanation — the report's What it does / Inputs / Outputs shape."""
        for term, body in pairs:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.6)
            r = p.add_run(term.replace("`", "").replace("**", "") + ": ")
            r.bold = True
            runs_into(p, body)

    def reference(self, n, text):
        """One bibliography entry: hanging indent, left-aligned so the numbers
        line up and the spacing is not stretched by justification."""
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent = Cm(1.2)
        pf.first_line_indent = Cm(-1.2)
        pf.space_after = Pt(7)
        r = p.add_run("[%d]" % n)
        r.bold = True
        p.add_run("\t")
        runs_into(p, text)
        return p

    def code(self, text):
        lines = text.strip("\n").rstrip().split("\n")
        for i, line in enumerate(lines):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(8 if i == 0 else 0)
            pf.space_after = Pt(8 if i == len(lines) - 1 else 0)
            pf.line_spacing = 1.0
            pf.left_indent = Cm(0.2)
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.keep_with_next = i < len(lines) - 1
            shade(p._p, CODE_SHADE)
            kinds = []
            if i == 0:
                kinds.append("top")
            if i == len(lines) - 1:
                kinds.append("bottom")
            kinds += ["left", "right"]
            borders(p, colour="D2D6DB", size="4", kinds=tuple(kinds))
            r = p.add_run(line if line.strip() else " ")
            r.font.name = CODE_FONT
            r.font.size = Pt(8.5)
            r.font.color.rgb = CODE_COLOR
        return p

    def note(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(10)
        shade(p._p, BOX_SHADE)
        borders(p, colour="D8C58A", size="6")
        runs_into(p, text, size=10.5)
        return p

    # ---- caption machinery ----
    def _caption(self, kind, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(14)
        p.paragraph_format.line_spacing = 1.0
        try:
            p.style = self.doc.styles["Caption"]
        except KeyError:
            pass
        r = p.add_run("%s " % kind)
        r.bold = True
        field(p, 'SEQ %s \\* ARABIC' % kind, placeholder="0")
        r = p.add_run(":  ")
        r.bold = True
        runs_into(p, text)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.name = BODY_FONT
            r.font.italic = False
            if r.font.color.rgb is None:
                r.font.color.rgb = RGBColor(0x33, 0x3A, 0x42)
        return p

    def figure(self, caption, image=None, width_cm=14.0, placeholder=None):
        """A real image if one exists, otherwise a clearly-marked drop-in box."""
        self.figures.append(caption)
        path = os.path.join(FIGS, image) if image else None
        if path and os.path.exists(path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            keep_with_next(p)
            p.add_run().add_picture(path, width=Cm(width_cm))
        else:
            body = placeholder or "Insert the final figure here."
            for i, line in enumerate(["[ FIGURE PLACEHOLDER ]", body]):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = p.paragraph_format
                pf.space_before = Pt(16 if i == 0 else 2)
                pf.space_after = Pt(16 if i else 2)
                pf.line_spacing = 1.0
                pf.keep_with_next = True
                shade(p._p, BOX_SHADE)
                kinds = ["left", "right"] + (["top"] if i == 0 else ["bottom"])
                borders(p, colour="C9A227", size="8", kinds=tuple(kinds))
                r = runs_into(p, line, size=10.5).runs[0]
                if i == 0:
                    for rr in p.runs:
                        rr.bold = True
                        rr.font.color.rgb = RGBColor(0x8A, 0x33, 0x00)
        return self._caption("Figure", caption)

    def table(self, caption, headers, rows, widths=None, size=9.5, align_right=()):
        self.tables.append(caption)
        cap = self._caption("Table", caption)
        # captions go above tables; move it before by inserting the table after it
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        hdr = t.rows[0]
        # Repeat the header on every continuation page, and never split a row
        # across a page break: a half-row with an empty first cell is unreadable.
        trpr = hdr._tr.get_or_add_trPr()
        trpr.append(_el("w:tblHeader", **{"w:val": "true"}))
        trpr.append(_el("w:cantSplit", **{"w:val": "true"}))
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = ""
            # The top-left cell belongs to the first COLUMN's shading, not the
            # header row's, so the row-label column reads as one continuous band.
            shade(cell._tc, FIRSTCOL_SHADE if i == 0 else HEADER_SHADE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(h.replace("`", "").replace("**", ""))
            r.bold = True
            r.font.size = Pt(size)
            r.font.name = BODY_FONT
        for row in rows:
            tr = t.add_row()
            tr._tr.get_or_add_trPr().append(_el("w:cantSplit", **{"w:val": "true"}))
            cells = tr.cells
            for i, val in enumerate(row):
                cell = cells[i]
                cell.text = ""
                if i == 0:
                    shade(cell._tc, FIRSTCOL_SHADE)
                p = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if i in align_right
                               else WD_ALIGN_PARAGRAPH.LEFT)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                runs_into(p, str(val), size=size)
                for r in p.runs:
                    if r.font.size is None:
                        r.font.size = Pt(size)
                    if r.font.name != CODE_FONT:
                        r.font.name = BODY_FONT
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)
        spacer.paragraph_format.line_spacing = 1.0
        return t
